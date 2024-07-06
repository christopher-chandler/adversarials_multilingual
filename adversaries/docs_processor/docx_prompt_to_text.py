# Standard
import platform

from io import BytesIO

# Pip
import pytesseract


from docx import Document
from PIL import Image


# Custom
from adversaries.docs_processor.tesseract import is_tesseract_installed

is_tesseract_installed()


class DocxPromptToText:
    """
    A class to extract text and images from a DOCX file.

    Args:
        docx_file (str): The path to the DOCX file.

    Attributes:
        docx_file (str): The path to the DOCX file.

    Methods:
        extract_images_from_docx(save_img=False) -> list:
            Extract images from the DOCX file.

        extract_text_from_docx() -> str:
            Extract text from the DOCX file.

        extract_text_from_docx_images() -> list:
            Extract text from images embedded in the DOCX file using OCR.

        docx_to_text(save_name: str) -> None:
            Write the contents of the DOCX file, including images, to a text file.
    """

    def __init__(self, docx_file):
        """
        Initialize the DocxPromptToText class.

        Args:
            docx_file (str): The path to the DOCX file.
        """
        self.docx_file = docx_file

    def extract_images_from_docx(self, save_img=False) -> list:
        """
        Extract images from the DOCX file.

        Args:
            save_img (bool, optional): Whether to save extracted images as files.
            Defaults to False.

        Returns:
            list: A list of PIL.Image objects representing the extracted images.
        """
        doc = Document(self.docx_file)
        images = []

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_data = BytesIO(rel.target_part.blob)
                img = Image.open(image_data)
                images.append(img)

        file, ext = self.docx_file.split(".")

        if save_img:
            for idx, img in enumerate(images):
                img.save(f"{file}_image_{idx}.png")  # or any other desired format

        return images

    def extract_text_from_docx(self) -> list:
        """
        Extract text from the DOCX file.

        Returns:
            str: The extracted text.
        """
        doc = Document(self.docx_file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)

    def extract_text_from_docx_images(self) -> list:
        """
        Extract text from images embedded in the DOCX file using OCR.

        Returns:
            list: A list of texts extracted from images.
        """
        docx_images = self.extract_images_from_docx()
        img_text = list()

        for idx, img in enumerate(docx_images):
            # Perform OCR using pytesseract

            text = pytesseract.image_to_string(img)
            img_text.append(text)

        return img_text

    def docx_to_text(self, save_name: str) -> None:
        """
        Args:
            save_name (str) save file for the prompt

        Write the contents of the docx including the images a text file.

        :return:
            None
        """

        docx_text = self.extract_text_from_docx()
        docx_img_text = self.extract_text_from_docx_images()

        with open(save_name, mode="w+") as out_file:
            out_file.write(docx_text)

            for image in docx_img_text:
                out_file.write(image)


if __name__ == "__main__":
    pass
